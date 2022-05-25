import os
import re
import chardet
import docx
from docx import Document
from sklearn.cluster import KMeans
import numpy as np
# 本代码的功能是将OCR的识别结果txt文件保存成save_path目录中的HXHGDCD6.docx
# 左页左栏段首区间最小值# 左页左栏段首区间最大值
l_l_min,l_l_max = 50,250 #120,230
# 左页右栏段首区间最小值# 左页右栏段首区间最大值
l_r_min,l_r_max = 940,1120#1010,1120
# 右页左栏段首区间最小值# 右页左栏段首区间最大值
r_l_min,r_l_max = 50,250 #120,230
# 右页左栏段首区间最小值# 右页右栏段首区间最大值
r_r_min,r_r_max = 960,1120#1010,1120
# 底边线 ，栏宽度
low_line,lan_w = 2690,780
# 保存的目录doc/及名字HXHGDCD6.docx
save_path = 'doc/HXHGDCD8.docx'
# OCR的识别结果txt文件 的 目录
dirs = "OCRresult/"
#去掉word中的括号和替换“--”为“-”，(因为正则提取时无法认识“--”)返回去掉括号的word
def deleKuo(word):

    l = word.find("(")
    r = word.find(")")
    ll = word.find("[")
    rr = word.find("]")
    n = 0
    while ll > -1 and rr > -1 and rr > ll:
        a = [word[0:ll], word[rr + 1:]]
        word = "".join(a)
        ll = word.find("[")
        rr = word.find("]")
        n += 1
        if n==2:
            break
    n = 0
    while l > -1 and r > -1 and r > l:
        a = [word[0:l], word[r + 1:]]
        word = "".join(a)
        l = word.find("(")
        r = word.find(")")
        n += 1
        if n==2:
            break
    word.replace("--","-")
    if len(word) == 0:
        return word
    if word[0] == "-":
        word = word[1:]
    return word

# 判断段首函数，输入word，输出True或者False
def is_para(word):
    word = word.replace(" ", '')
    ps = re.search("^(([\s\w，]{1,8}-){0,1}[\u4E00-\u9FA5·\s]{2,20}-{0,1}){1,4}[ⅠⅡⅢⅣⅤA-Z\d\s]{0,6}"
                   "(([\s\w，]{1,8}-){0,1}[A-Za-z]{1}[A-Za-z\s\d；]{4,40}-{0,1}){1,4}[ⅠⅡⅢⅣⅤA-Z\d\s]{0,6}"
                   , word)
    if ps:
        print("word:", word)
        print('pppppppps:', ps.group())
        return True
    p2 = re.search("^(([\s\w，]{1,8}-){0,1}[\u4E00-\u9FA5·\s]{1,20}-{0,1}){1,4}[A-Z\d\s]{0,6}"
                   "见[\u4E00-\u9FA5]{0,20}([(][\u4E00-\u9FA5][)]){0,1}[\u4E00-\u9FA5]+\d{1,5}"
                   , word)
    if p2:
        print("word:", word)
        print('pppppppp2:', p2.group())
        return True
    # 122_苯乙酸对甲酚酯p-cre syl phenyl acetate
    # 苯基膦酸丁酯butyl phenyl phosphonate； di but yphenyl phosphonate
    # 112(-) - 1 - 苯基乙胺(一) - 1 - phenylethylamine
    # p5 = re.search("^([\d，]{0,8}-{0,1}[()-\u4E00-\u9FA5\s]{2,20}-{0,1}"
    #                "[\d，]{0,8}-{0,1}[a-z\s]{5,40}-{0,1}){1,3}"
    #                , word)
    # 2，3-苯·酚[a]苯H3F
    # DX编码DX code
    # 124_Kinyon泵Kinyon pump
    # Worthington泵Worthington pump
    p3 = re.search("^(([A-Z][a-z\s]{3,40})|([A-Z]{1,10}))[\u4E00-\u9FA5]{1,20}\s?[A-Z][a-z\s]{3,40}"
                   , word)
    if p3:
        print("word:", word)
        print('pppppppp3:', p3.group())
        return True
    # 4-(2-吡啶偶氮)间苯二酚4-(2-pyridyl azo) resorcinol；PAR 未解决
    # 1 - (吡啶基偶氮) - 2 - 萘酚1 - (2 - pyridyl azo) - 2 - naphthol；
    # p4 = re.search("^([\s\w，-]{0,9}[\u4E00-\u9FA5\s]{2,10}|([(][\s\d，-]{0,9}[\u4E00-\u9FA5]{2,20}[)][\u4E00-\u9FA5\s]{2,10})-{0,1}){1,4}"
    #                "([\s\w，-]{0,9}[a-z\s]{2,20}|([(][\s\w，-]{0,9}[a-z\s；]{2,20}[)])-{0,1}){1,4}"
    #                , word)
    # 123_苯乙烯-丁二烯-苯乙烯嵌段共聚物/高抗冲聚苯乙烯共混物styrene-butadiene-styrene block copolymer/high impacpolystyrene blend；
    # 苯乙烯-丁二烯(共聚物) 塑料styrene-butadiene(co polymer) plastics
    # 乙烯基(甲) 酮phenyl vinyl ket on
    # 苯甲酸(2， 2， 6， 6-四甲基哌啶醇酯) 2， 2， 6， 6-tetramethyl piperidine benzoate
    # 117_0-(苯硫基甲基) 羟胺O-(phenyl thio methyl) hydroxylamineC6H， SCH2ONH 2
    # 苯基(三氟甲基) 汞phenyl(trifluoromethyl) mercuryC6HsHgCF 3
    # 1， 3-苯二甲酰-1， 1-双(2-甲基氮丙啶) 1， 3-phenylene di-carbonyl-1， 1-bis(2-metHsC-C-NcHs
    # 苯基(二溴氯甲基) 汞phenyl(dibromo chloromethyl) mercuryCfHsHgCBr2Cl
    # 苯(基)甲基硅油phenyl methyl silicone fluid
    word = deleKuo(word)#此处去括号
    p111 = re.search("^(([\w，]{1,8}-){0,2}[\u4E00-\u9FA5/]{1,20}-{0,1}){1,4}"
                   "(([\w，]{1,8}-){0,2}[A-Za-z][A-Za-z\d；/]{4,40}-{0,1}){1,4}"
                   , word)
    if p111:
        print("word:", word)
        print('pppppppp111:', p111.group())
        return True
    return False

#K-mean聚类，可以把点集分为k类,输入需要分类的点centers和类别数k，输出最小缩进距离min_indent
def k_means(centers,k):
    np_centers = np.array(centers).reshape(-1,1)
    # 训练模型
    model = KMeans(n_clusters=k)
    model.fit(np_centers)
    # 分类中心点坐标
    np_centers = model.cluster_centers_
    # 预测结果
    result = model.predict(np_centers)
    print(centers)
    print(model.labels_)
    min_indent0,min_indent1,min_indent2 = 0,0,0
    for i,ind in enumerate(model.labels_):
        if ind == 0 and centers[i] > min_indent0:
            min_indent0 = centers[i]
        if ind == 1 and centers[i] > min_indent1:
            min_indent1 = centers[i]
        if ind == 2 and centers[i] > min_indent2:
            min_indent2 = centers[i]
    min_indent = min_indent0
    if min_indent1 < min_indent0:
        min_indent = min_indent1
    if min_indent2 < min_indent1 and min_indent2 < min_indent0:
        min_indent = min_indent2
    return min_indent

# 收集一栏的文字this_lan = ["",""]和位置this_pos = [150,160]
def lan_pos(prism_wordsInfo,min,max):
    this_lan = []
    this_pos = []
    for prism in prism_wordsInfo:
        poss = prism["pos"]
        pos = poss[0]['x']
        y = poss[0]['y']
        # print("pos",pos)
        # 去除x:min以下，x:max+750以上的字 和 小于low_line的字
        if pos < min or pos > max + lan_w or y > low_line:
            continue
        this_pos.append(pos)
        word = prism["word"]
        this_lan.append(word)
    return this_lan,this_pos

# 初选出这一栏的位于行首的字,存在
def first_select_indent(this_lan,this_pos,min,max):
    first_select = []
    for i, cw in enumerate(this_lan):
        pos = this_pos[i]
        if pos > min and pos < max:
            first_select.append(pos)
    return first_select

# 缩进一栏函数，利用K-mean聚类得到的阈值min_indent和区间值max和段首语义is_para()判断是否需要缩进
def indent(prism_wordsInfo,min,max,doc,paraObj_l,page_num):
    # average,l = avg(prism_wordsInfo,min,max)
    # 收集一栏的所有文字this_lan = ["",""]和位置this_pos = [150,160]
    this_lan, this_pos = lan_pos(prism_wordsInfo,min,max)
    l = len(this_lan)
    #print('l:',l,"p:",len(this_pos))
    first_select = first_select_indent(this_lan,this_pos,min,max)
    min_indent = k_means(first_select,3)
    #print("min_indent:",min_indent)
    for i,cw in enumerate(this_lan):
        if i<l-2:
            w = cw + this_lan[i+1] + this_lan[i+2]
        elif i ==l-2:
            w = cw + this_lan[i + 1]
        else:
            w = cw
        pos = this_pos[i]
        if pos >= min_indent and pos < max and is_para(w):
            # 写段落
            print(page_num,pos,"####",cw)
            paraObj_l = doc.add_paragraph(str(page_num) + '_' + cw)
        else :
            #续写段落
            print(page_num,pos,cw)
            paraObj_l.add_run(cw)
    print("indent over!")
    return paraObj_l

#从图片名字拿到它的页码,和它的页面类型，1为左2为右
def num_from_name(filename):
    num = re.search("\d+",filename)
    print("num",num)
    if num :
        num = num.group()
        return int(num)
    else:
        print("名字有误：",num,"页面类型有误")

# 从图片中拿到它页码和辨别左右页
def num_lr(prism_wordsInfo,pic_name):
    lr = ""
    num = -1
    for prism in prism_wordsInfo:
        poss = prism["pos"]
        x = poss[0]['x']
        y = poss[0]["y"]
        #print("x,y:",x,y)
        if y < low_line:
            continue
        word = prism["word"]
        print("x,y:", x, y)
        if  x > 150 and x < 800:  # 左单页
            lr = "左"
            n = re.findall("\d+",word)
            if len(n)>0:
                num = n[0]
                print("num:",num)
        if lr == "左"  and x > 3000:#双页
            lr = "双"
        if  x>1500 and x<2000:#右单页
            lr = "右"
            print("youword:",word)
            n = re.findall("\d+", word)
            if len(n)>0:
                num = n[0]
                print("num:",num)
    if num == -1:
        print("未检测页码，请手动输入页码！！！！！！！",pic_name)
        num = int(input("未检测页码，请手动输入页码"))
        # if num>0:
        #     num = int(num)
        # else:
        #     while num == :
        #         num = input("未检测页码，请手动输入页码")
        #         num = int(num)
    if lr == '':
        print("未检测页码，请手动输入 左 、右、或者 双",pic_name)
        lr = input("未检测页码，请手动输入 左 、右、或者 双")
        # while lr not in ["左","右","双"]:
        #     page = input("未检测页码，请手动输入 左 、右、或者 双")
    return num,lr

#利用以上函数，将文字写入doc中，并排版
def write_doc(pic_name,doc,paraObj_l):
    with open('OCRresult/' + pic_name , 'rb') as f:
        data = f.read()
    type = chardet.detect(data)
    result = data.decode(type["encoding"])
    page_dict = eval(result)
    prism_wordsInfo = page_dict["prism_wordsInfo"]
    #page_num,page_type = num_lr(prism_wordsInfo,pic_name)
    page_num = num_from_name(pic_name)
    print("page_num:",page_num)
    # 1:左  2:右
    if page_num % 2 == 0:
        page_type = 1
    else:
        page_type = 2
    # 左页排版
    if page_type == 1:
        paraObj_l = indent(prism_wordsInfo, l_l_min, l_l_max,doc,paraObj_l,page_num)
        paraObj_l = indent(prism_wordsInfo, l_r_min, l_r_max,doc,paraObj_l,page_num)
    # 右叶排版
    if page_type == 2:
        paraObj_l = indent(prism_wordsInfo, r_l_min, r_l_max,doc,paraObj_l,page_num)
        paraObj_l = indent(prism_wordsInfo, r_r_min, r_r_max,doc,paraObj_l,page_num)
    return paraObj_l

if __name__ == '__main__':
    doc = Document()
    paraObj_l = doc.add_paragraph('')
    #dirs = "OCRresult/"
    for dir in os.listdir(dirs):
        print("dirs:",dir)
        # 写入到docx文件中
        paraObj_l = write_doc(dir,doc,paraObj_l)
    doc.save(save_path)


