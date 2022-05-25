import docx
from docx import Document
import re
import csv
from xml.dom.minidom import *
from parse_utils import r1,r2,r3
# 本代码的功能是将docx文件转换为xml文件
# docx的源目录为docx_dir,保存的目录路径为save_path_xml
docx_dir = 'doc/HXHGDCD8.docx'
save_path_xml = 'xml/HXHGDCD8.xml'

#创建一个文档对象
doc=xml.dom.minidom.Document()
#创建一个根节点
root=doc.createElement('contents')
#根节点加入到tree
doc.appendChild(root)
#创建二级节点，并连接
source=doc.createElement('source')
entries=doc.createElement('entries')
root.appendChild(source)
root.appendChild(entries)
source.appendChild(doc.createTextNode('HXHGDCD')) #添加文本节点
#去掉字符中的括号和替换“--”为“-”，因为正则提取时无法认识“--”
def deleKuo(word):
    l = word.find("(")
    r = word.find(")")
    ll = word.find("[")
    rr = word.find("]")
    n = 0
    while ll > -1 and rr > -1 and rr > ll:
        a = [word[0:ll], word[rr + 1:]]
        word = "".join(a)
        ll = word.find("[")
        rr = word.find("]")
        n += 1
        if n==2:
            break
    n = 0
    while l > -1 and r > -1 and r > l:
        a = [word[0:l], word[r + 1:]]
        word = "".join(a)
        l = word.find("(")
        r = word.find(")")
        n += 1
        if n==2:
            break
    word.replace("--","-")
    return word
#将以下值写入xml：页码值             名字         同义英文值          同意中文值      内容值          另外名字
def write_entry(pageNumber_value ,name_value ,synonym_en_value ,synonym_value ,content_value ,another_name):

    entry = doc.createElement('entry')
    entries.appendChild(entry)
    # 第四级节点
    pageNumber = doc.createElement('pageNumber')
    entry.appendChild(pageNumber)
    pageNumber.appendChild(doc.createTextNode(pageNumber_value)) #写入内容

    name = doc.createElement('name')
    entry.appendChild(name)
    name.appendChild(doc.createTextNode(name_value))  # 写入内容

    #英文不为空时新建标签
    if synonym_en_value:
        synonym_en = doc.createElement('synonym')
        synonym_en.appendChild(doc.createTextNode(synonym_en_value))  # 写入内容
        synonym_en.setAttribute('language', 'en')
        entry.appendChild(synonym_en)
    # 别名不为空时新建标签
    if synonym_value:
        for i in range(len(synonym_value)):
            synonym = doc.createElement('synonym')
            synonym.appendChild(doc.createTextNode(synonym_value[i]))  # 写入内容
            entry.appendChild(synonym)
    #
    if another_name:
        for another_nameX in another_name:
            synonym = doc.createElement('synonym')
            synonym.appendChild(doc.createTextNode(another_nameX))  # 写入内容
            entry.appendChild(synonym)
    # 内容不为空时新建标签
    if content_value:
        content = doc.createElement('content')
        content.appendChild(doc.createTextNode(content_value))  # 写入内容
        entry.appendChild(content)
# 提取描述(除了中英文名字)
def extract_describe(paragraph_text,term_and_ename):
    #替换段落paragraph_text中的term_and_ename词语为空格，去调首位空格并返回
    describe = paragraph_text.replace(term_and_ename, '')
    describe = describe.strip()
    return describe
# 设置页码值pageNumber_value名字name_value同义英文值synonym_en_value同意中文值synonym_value内容值content_value另外名字another_name
def set_val(paragraph,data = {},pageNumber = "-1",term = '',synonym_en = '',synonym = [],describe = '',another_name = []):
    #第一参数：页码pageNumber
    pageNumber_tem = re.search("^(\d{1,5})_",paragraph)
    if pageNumber_tem:
        pageNumber = pageNumber_tem.group(1)
        #print("pagenum",pageNumber)
    #第二个参数，术词term
    term,synonym_en,synonym,ll = r2(paragraph)
    if term == None:
        term,synonym_en,synonym,ll = r1(paragraph)
        print("r1:", term, "  ",synonym_en,"  ", synonym)
    if term == None:
        term,synonym_en,synonym,ll = r3(paragraph)
        print("r3:", term,"  ", synonym_en,"  ", synonym)
    print(paragraph)
    describe = extract_describe(paragraph,paragraph[0:ll])
    print(describe)
    data["pageNumber"] = pageNumber
    data["term"] = term
    data["synonym_en"] = synonym_en
    data["synonym"] = synonym
    data["describe"] = describe
    data["another_name"] = another_name
    # print(data)
    return data

if __name__ == '__main__':
    document = docx.Document(docx_dir)
    for paragraph in document.paragraphs:
        data = set_val(paragraph.text)
        if data["term"]!=None:
            write_entry(data["pageNumber"], data["term"], data["synonym_en"]
                        , data["synonym"], data["describe"],data["another_name"])
    with open(save_path_xml,'w',encoding='utf-8') as f:
        doc.writexml(f,indent='',addindent='\t',newl='\n',encoding='utf-8')


# def saveData(csv_obj,data):
#     #在csv_obj每行写入data
#     csv.writer(csv_obj).writerow(data)